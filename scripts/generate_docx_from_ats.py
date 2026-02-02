import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _set_default_font(document: Document, name: str = "Calibri", size_pt: int = 11) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = name
    font.size = Pt(size_pt)


def _add_bullet_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    _set_default_font(doc)

    for raw in md:
        line = raw.rstrip()
        if not line.strip():
            continue

        # Name
        if line.startswith("## "):
            heading = line[3:].strip()
            # Treat the very first H2 as title; others as section headers.
            if not any(p.text.strip() for p in doc.paragraphs):
                p = doc.add_paragraph()
                run = p.add_run(heading)
                run.bold = True
                run.font.size = Pt(16)
            else:
                doc.add_heading(heading, level=1)
            continue

        # Section heading
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue

        # Bold-only line like **Role**
        m = re.fullmatch(r"\*\*(.+?)\*\*", line.strip())
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1).strip())
            run.bold = True
            continue

        # Bullets
        if line.lstrip().startswith("❖"):
            _add_bullet_paragraph(doc, line.lstrip().lstrip("❖").strip())
            continue

        # Soft line breaks and markdown artifacts
        cleaned = line.replace("  ", " ").strip()
        doc.add_paragraph(cleaned)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md_path = root / "cv" / "Mishab_Abdurahman_ATS_Resume.md"
    docx_path = root / "cv" / "Mishab_Abdurahman_ATS_Resume.docx"

    if not md_path.exists():
        raise SystemExit(f"Missing input file: {md_path}")

    md_to_docx(md_path, docx_path)
    print(f"Wrote: {docx_path}")
